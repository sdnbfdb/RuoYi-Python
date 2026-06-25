from flask import Flask, send_from_directory
from flask_cors import CORS
from flask import jsonify, request
import os
import uuid
import requests
from datetime import datetime
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from layer.user.user import register_user_routes
from knowledge.knowledge import (
    get_knowledge_list,
    get_knowledge_detail,
    create_knowledge,
    update_knowledge,
    delete_knowledge,
    upload_attachment,
    serve_attachment,
    get_statistics,
    load_knowledge_from_db
)
from layer.user.organization import OrganizationManager
from old.api_call import call_api, call_api_with_tools
from old.phone.link import (
    send_to_connected_user,
    register_connected_user,
    disconnect_user,
    get_connected_users,
    update_heartbeat,
    check_expired_users,
    get_unread_messages
)
from routes.chat_routes import (
    _parse_table_action,
    _format_table_result,
    _parse_file_content
)
from old.tool.tool import (
    call_knowledge_base,
    call_organization,
    call_search_api,
    generate_image,
    get_task_status,
    call_table_tool
)
import socket

# 会话表格数据存储
session_table_data = {}

# 缓存局域网IP，避免每次请求都建立socket连接
cached_local_ip = None

def get_local_ip():
    """获取本机局域网IP并缓存"""
    global cached_local_ip
    if cached_local_ip:
        return cached_local_ip
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.settimeout(0.3)
    try:
        s.connect(('8.8.8.8', 80))
        cached_local_ip = s.getsockname()[0]
    except Exception:
        cached_local_ip = '127.0.0.1'
    finally:
        s.close()
    return cached_local_ip

app = Flask(__name__)

IMAGE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'image')

# 组织架构管理器实例
org_manager = OrganizationManager()

def enhance_prompt(message):
    """根据@前缀增强prompt，插入真实后端数据"""
    enhanced = message
    
    if '@知识库' in message:
        try:
            data = load_knowledge_from_db()
            stats = {
                'total': len(data),
                'by_type': {}
            }
            for item in data:
                t = item.get('type', 'other')
                stats['by_type'][t] = stats['by_type'].get(t, 0) + 1
            titles = [item['title'] for item in data[:15]]
            context = f"【知识库真实数据】共有{stats['total']}条记录"
            if stats['by_type']:
                type_desc = ', '.join([f"{k}:{v}条" for k, v in list(stats['by_type'].items())[:5]])
                context += f"，类型分布：{type_desc}"
            context += f"。主要条目包括：{', '.join(titles[:10])}"
            if len(titles) > 10:
                context += f"等共{len(titles)}条"
            context += "。请基于以上真实数据回答，不要编造。"
            enhanced = enhanced + '\n\n' + context
        except Exception as e:
            print(f"[enhance_prompt] 获取知识库数据失败: {e}")
    
    if '@组织管理' in message or '@组织架构' in message:
        try:
            tree = org_manager.build_org_tree()
            def count_depts(node):
                cnt = 1
                for c in node.get('children', []):
                    cnt += count_depts(c)
                return cnt
            total = count_depts(tree)
            top_names = [c.get('name', '') for c in tree.get('children', [])[:10]]
            context = f"【组织架构真实数据】共有{total}个部门"
            if top_names:
                context += f"，一级部门包括：{', '.join(top_names)}"
            context += "。请基于以上真实数据回答，不要编造。"
            enhanced = enhanced + '\n\n' + context
        except Exception as e:
            print(f"[enhance_prompt] 获取组织架构数据失败: {e}")
    
    return enhanced


# ========== Function Calling 工具定义 ==========

TOOLS_KNOWLEDGE = [
    {
        "type": "function",
        "function": {
            "name": "query_knowledge",
            "description": "查询煤矿知识库，获取煤矿安全、技术、管理等方面的资料。支持搜索、列表、详情、统计、向量搜索等操作。",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["list", "detail", "search", "extract", "stats", "vector_search"],
                        "description": "操作类型：list(列表), detail(详情), search(搜索), extract(提取), stats(统计), vector_search(向量搜索)"
                    },
                    "search": {
                        "type": "string",
                        "description": "搜索关键词，当action为search或vector_search时使用"
                    },
                    "knowledge_id": {
                        "type": "string",
                        "description": "资料ID，当action为detail或extract时使用"
                    },
                    "top_k": {
                        "type": "integer",
                        "description": "向量搜索返回结果数量"
                    }
                },
                "required": ["action"]
            }
        }
    }
]

TOOLS_ORGANIZATION = [
    {
        "type": "function",
        "function": {
            "name": "query_organization",
            "description": "查询煤矿组织架构信息，包括部门层级、负责人、上下级关系等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["get", "list"],
                        "description": "操作类型：get(查询单个), list(列表)"
                    },
                    "name": {
                        "type": "string",
                        "description": "部门名称或ID"
                    },
                    "parent_id": {
                        "type": "string",
                        "description": "父部门ID或名称，用于list过滤"
                    }
                },
                "required": ["action"]
            }
        }
    }
]

TOOLS_SEARCH = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "使用博查搜索API进行联网搜索，获取最新的煤矿行业新闻、政策、技术等信息。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索关键词"
                    }
                },
                "required": ["query"]
            }
        }
    }
]

TOOLS_PERSONAL_ASSISTANT = [
    {
        "type": "function",
        "function": {
            "name": "query_knowledge",
            "description": "查询煤矿知识库，获取煤矿安全、技术、管理等方面的资料。支持搜索、列表、详情、统计、向量搜索等操作。",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {"type": "string", "enum": ["list", "detail", "search", "extract", "stats", "vector_search"]},
                    "search": {"type": "string", "description": "搜索关键词"},
                    "knowledge_id": {"type": "string", "description": "资料ID"},
                    "top_k": {"type": "integer", "description": "向量搜索返回结果数量"}
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "query_organization",
            "description": "查询煤矿组织架构信息。",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {"type": "string", "enum": ["get", "list"]},
                    "name": {"type": "string", "description": "部门名称或ID"},
                    "parent_id": {"type": "string", "description": "父部门ID"}
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "联网搜索，获取最新的煤矿行业新闻、政策、技术等信息。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "搜索关键词"}
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "使用通义万相AI生成图片。当用户要求生成带图文案、配图、插图、海报、宣传图等时，必须调用此工具生成与文案内容主题匹配的AI图片。支持根据文本描述生成任意主题的图片。",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "图片描述文本，建议用英文描述，包含画面主体、风格、场景等细节"},
                    "size": {"type": "string", "enum": ["1024*1024", "720*1280", "1280*720"], "description": "图片尺寸，小红书配图建议 720*1280"}
                },
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "table_operation",
            "description": "表格操作工具，支持创建表格、查看表格、添加行列、删除行列、生成图表、导出表格等操作。",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {"type": "string", "enum": ["table_create", "table_delete", "table_add_row", "table_delete_row", "table_add_column", "table_delete_column", "table_statistics", "table_chart", "table_export", "table_filter"]},
                    "headers": {"type": "array", "items": {"type": "string"}, "description": "表头列表"},
                    "column_name": {"type": "string", "description": "列名"},
                    "row_index": {"type": "integer", "description": "行索引"},
                    "chart_type": {"type": "string", "enum": ["bar", "line", "pie", "scatter", "histogram"], "description": "图表类型"},
                    "export_format": {"type": "string", "enum": ["csv", "json", "markdown"], "description": "导出格式"}
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "clean_text",
            "description": "清理文本，去除噪声、HTML标签、数字、标点等。用于预处理用户输入的文本数据。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "需要清理的原始文本"},
                    "remove_num": {"type": "boolean", "description": "是否去除数字"},
                    "remove_punc": {"type": "boolean", "description": "是否去除标点"},
                    "remove_html": {"type": "boolean", "description": "是否去除HTML标签"}
                },
                "required": ["text"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "extract_text",
            "description": "从文本中提取特定类型的内容（中文、数字、邮箱、手机号、网址）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "需要提取内容的文本"},
                    "extract_type": {"type": "string", "enum": ["chinese", "numbers", "emails", "phones", "urls"], "description": "提取类型"}
                },
                "required": ["text", "extract_type"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "split_text",
            "description": "将长文本按Markdown标题和句号分块，支持自定义分块大小和重叠。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "需要分块的文本"},
                    "max_size": {"type": "integer", "description": "每个分块的最大字符数"},
                    "overlap": {"type": "integer", "description": "分块之间的重叠字符数"}
                },
                "required": ["text"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "analyze_table_relationships",
            "description": "分析多个表格文件之间的表头关联关系，查找共享表头和文件间关系。",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory_path": {"type": "string", "description": "包含表格文件的目录路径"}
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "retrieve_table_data",
            "description": "根据表头匹配检索表格数据，支持多表连接和条件过滤。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query_header": {"type": "string", "description": "查询表头名称"},
                    "top_n": {"type": "integer", "description": "返回的最大记录数"},
                    "directory_path": {"type": "string", "description": "包含表格文件的目录路径"}
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "find_similar_headers",
            "description": "查找与查询表头相似的表头，用于表格关联分析。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query_header": {"type": "string", "description": "查询表头名称"},
                    "top_k": {"type": "integer", "description": "返回的相似表头数量"},
                    "directory_path": {"type": "string", "description": "包含表格文件的目录路径"}
                },
                "required": ["query_header"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_header_similarity",
            "description": "计算两个表头之间的相似度，用于判断列是否可以连接。",
            "parameters": {
                "type": "object",
                "properties": {
                    "header1": {"type": "string", "description": "第一个表头名称"},
                    "header2": {"type": "string", "description": "第二个表头名称"}
                },
                "required": ["header1", "header2"]
            }
        }
    }
]


def _query_knowledge_handler(action='search', search='', knowledge_id='', **kwargs):
    """知识库工具执行处理"""
    if action == 'search':
        return call_knowledge_base(action, search=search or kwargs.get('keyword', ''))
    elif action == 'detail':
        return call_knowledge_base(action, knowledge_id=knowledge_id)
    elif action == 'list':
        return call_knowledge_base(action)
    elif action == 'extract':
        return call_knowledge_base(action, knowledge_id=knowledge_id)
    elif action == 'stats':
        return {"success": True, "data": {"message": get_statistics().get_data(as_text=True)}}
    elif action == 'vector_search':
        top_k = kwargs.get('top_k', 5)
        try:
            from vector_store.chroma_store import CharmDBStore
            chroma_store = CharmDBStore(persist_directory=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chroma_db'))
            results = chroma_store.similarity_search_with_score(
                collection_name='knowledge_base',
                query=search,
                k=top_k
            )
            if results:
                result_str = f"向量搜索 '{search}' 找到 {len(results)} 条记录：\n"
                for i, (doc, score) in enumerate(results, 1):
                    metadata = doc.metadata
                    result_str += f"{i}. {metadata.get('title', '')} - 章节: {metadata.get('section_title', '')} (相似度: {score:.4f})\n"
                return {"success": True, "data": {"message": result_str}}
            else:
                return {"success": True, "data": {"message": f"向量搜索未找到与 '{search}' 相关的记录"}}
        except Exception as e:
            return {"success": True, "data": {"message": f"向量搜索失败: {str(e)}"}}
    else:
        return call_knowledge_base(action, **kwargs)


def _query_organization_handler(action='list', name='', parent_id='', **kwargs):
    """组织架构工具执行处理"""
    if action in ('get', 'get_by_name'):
        return call_organization('get', name=name or kwargs.get('org_id', ''))
    elif action == 'list':
        return call_organization('list', parent=parent_id or kwargs.get('parent', ''))
    else:
        return call_organization(action, **kwargs)


def _web_search_handler(query=''):
    """联网搜索工具执行处理"""
    return call_search_api(query)


def _generate_image_handler(prompt='', size='1024*1024', **kwargs):
    """图片生成工具执行处理（支持异步任务轮询）"""
    result = generate_image(prompt, size=size, **kwargs)
    if result.get('success') and result.get('data', {}).get('async'):
        # 异步任务，轮询获取结果
        task_id = result['data']['task_id']
        import time
        for _ in range(30):  # 最多轮询30次，每次2秒
            time.sleep(2)
            status = get_task_status(task_id, task_type='image')
            if status.get('success'):
                output = status.get('data', {})
                if output.get('task_status') == 'SUCCEEDED':
                    return status
                elif output.get('task_status') == 'FAILED':
                    return {"success": False, "message": "图片生成失败", "data": output}
        return {"success": False, "message": "图片生成超时，请稍后查询任务状态", "task_id": task_id}
    return result


def _table_operation_handler(action='', **kwargs):
    """表格操作工具执行处理"""
    from old.tool.tool import call_table_tool
    return call_table_tool(action, **kwargs)


# 工具执行映射
TOOL_HANDLERS_KNOWLEDGE = {
    'query_knowledge': _query_knowledge_handler
}

TOOL_HANDLERS_ORGANIZATION = {
    'query_organization': _query_organization_handler
}

TOOL_HANDLERS_SEARCH = {
    'web_search': _web_search_handler
}

def _clean_text_handler(text='', remove_num=False, remove_punc=False, remove_html=True):
    """文本清理工具处理"""
    from nlp.clear import TextCleaner
    cleaner = TextCleaner()
    cleaned = cleaner.clean_text(text, remove_num=remove_num, remove_punc=remove_punc, remove_html=remove_html)
    return {"success": True, "data": {"message": f"清理后的文本（长度：{len(cleaned)}字符）:\n{cleaned[:2000]}{'...' if len(cleaned) > 2000 else ''}"}}


def _extract_text_handler(text='', extract_type='chinese'):
    """文本提取工具处理"""
    from nlp.clear import TextCleaner
    import re
    cleaner = TextCleaner()
    
    if extract_type == "chinese":
        result = cleaner.extract_chinese(text)
        msg = f"提取的中文内容（长度：{len(result)}字符）:\n{result[:2000]}{'...' if len(result) > 2000 else ''}"
    elif extract_type == "numbers":
        result = cleaner.extract_numbers(text)
        msg = f"提取的数字列表（共{len(result)}个）:\n{', '.join(result[:50])}{'...' if len(result) > 50 else ''}"
    elif extract_type == "emails":
        emails = re.findall(r'[\w.-]+@[\w.-]+\.\w+', text)
        msg = f"提取的邮箱列表（共{len(emails)}个）:\n{', '.join(emails[:20])}{'...' if len(emails) > 20 else ''}"
    elif extract_type == "phones":
        phones = re.findall(r'1[3-9]\d{9}', text)
        msg = f"提取的手机号列表（共{len(phones)}个）:\n{', '.join(phones[:20])}{'...' if len(phones) > 20 else ''}"
    elif extract_type == "urls":
        urls = re.findall(r'https?://[\w\-._~:/?#[\]@!$&\'()*+,;=%]+', text)
        msg = f"提取的URL列表（共{len(urls)}个）:\n{', '.join(urls[:20])}{'...' if len(urls) > 20 else ''}"
    else:
        msg = f"不支持的提取类型: {extract_type}"
    return {"success": True, "data": {"message": msg}}


def _split_text_handler(text='', max_size=800, overlap=100):
    """文本分块工具处理"""
    import re
    lines = text.split('\n')
    sections = []
    current_title = '概述'
    current_lines = []
    
    for line in lines:
        stripped = line.strip()
        h2_match = re.match(r'^##\s+(.+)', stripped)
        if h2_match:
            if current_lines:
                content = '\n'.join(current_lines).strip()
                if content:
                    sections.append({'section_title': current_title, 'content': content})
            current_title = h2_match.group(1).strip()
            current_lines = [stripped]
        elif stripped.startswith('# '):
            if current_lines:
                content = '\n'.join(current_lines).strip()
                if content:
                    sections.append({'section_title': current_title, 'content': content})
            current_title = stripped.lstrip('#').strip()
            current_lines = [stripped]
        else:
            current_lines.append(line)
    
    if current_lines:
        content = '\n'.join(current_lines).strip()
        if content:
            sections.append({'section_title': current_title, 'content': content})
    
    if not sections:
        sections.append({'section_title': '全文', 'content': text.strip()})
    
    chunks = []
    for section in sections:
        content = section['content']
        if len(content) <= max_size:
            chunks.append({'section_title': section['section_title'], 'content': content})
        else:
            sentences = re.split(r'(?<=[。！？；])', content)
            sentences = [s for s in sentences if s.strip()]
            current_chunk = ''
            for sent in sentences:
                if len(current_chunk) + len(sent) > max_size and current_chunk:
                    chunks.append({'section_title': section['section_title'], 'content': current_chunk.strip()})
                    if overlap > 0 and len(current_chunk) > overlap:
                        current_chunk = current_chunk[-overlap:] + sent
                    else:
                        current_chunk = sent
                else:
                    current_chunk += sent
            if current_chunk.strip():
                chunks.append({'section_title': section['section_title'], 'content': current_chunk.strip()})
    
    result_lines = [f"文本已分成 {len(chunks)} 个分块："]
    for i, chunk in enumerate(chunks):
        preview = chunk['content'][:100] + '...' if len(chunk['content']) > 100 else chunk['content']
        result_lines.append(f"\n分块 {i+1}（章节: {chunk['section_title']}）:\n  长度: {len(chunk['content'])} 字符\n  预览: {preview}")
    
    return {"success": True, "data": {"message": '\n'.join(result_lines)}}


def _analyze_table_relationships_handler(directory_path=''):
    """表格关联分析工具处理"""
    try:
        from nlp.excel.header_matrix import HeaderMatrixAnalyzer
        import os
        
        if not directory_path:
            directory_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
        
        if not os.path.exists(directory_path):
            return {"success": True, "data": {"message": f"目录不存在: {directory_path}"}}
        
        analyzer = HeaderMatrixAnalyzer()
        analyzer.load_files_from_directory(directory_path)
        analyzer.build_association_matrix()
        
        result_lines = [f"表头关联分析结果:\n  加载文件数: {len(analyzer.file_list)}\n  不同表头数: {len(analyzer.all_headers)}"]
        
        common_headers = analyzer.find_common_headers(min_files=2)
        if common_headers:
            result_lines.append(f"\n  公共表头（出现在多个文件中）:")
            for header, count in common_headers[:5]:
                result_lines.append(f"    - {header}: 出现在 {count} 个文件中")
        
        relationships = analyzer.find_file_relationships()
        if relationships:
            result_lines.append(f"\n  文件间关系（共享表头）:")
            for rel in relationships[:5]:
                result_lines.append(f"    - {rel['file1']} ↔ {rel['file2']}（共享{rel['shared_headers']}个表头）")
        
        return {"success": True, "data": {"message": '\n'.join(result_lines)}}
    except Exception as e:
        print(f'[_analyze_table_relationships_handler] 错误: {e}', flush=True)
        return {"success": True, "data": {"message": f"表格关联分析失败: {str(e)}"}}


def _retrieve_table_data_handler(query_header='', top_n=50, directory_path=''):
    """表格数据检索工具处理"""
    try:
        from nlp.excel.table_retriever import TableDataRetriever
        from nlp.excel.header_embedding import HeaderEmbeddingEncoder
        import os
        
        if not directory_path:
            directory_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
        
        if not os.path.exists(directory_path):
            return {"success": True, "data": {"message": f"目录不存在: {directory_path}"}}
        
        retriever = TableDataRetriever(data_dir=directory_path)
        retriever.load_all()
        
        if not retriever.tables:
            return {"success": True, "data": {"message": f"目录中未找到CSV文件: {directory_path}"}}
        
        if query_header:
            encoder = HeaderEmbeddingEncoder()
            encoder.encode_directory(directory_path)
            matched_headers = encoder.find_similar_headers(query_header, top_k=10)
        else:
            matched_headers = []
        
        result = retriever.retrieve(matched_headers=matched_headers, top_n=top_n)
        
        result_lines = [f"表格检索结果:\n  加载文件: {', '.join(result['files_loaded'])}\n  总记录数: {result['total']}\n  返回记录: {len(result['records'])}\n  列: {', '.join(result['headers'])}"]
        
        if result['records']:
            result_lines.append(f"\n  示例数据:")
            for i, row in enumerate(result['records'][:5]):
                row_str = ", ".join(f"{k}: {v}" for k, v in row.items())
                result_lines.append(f"    {i+1}. {row_str}")
        
        return {"success": True, "data": {"message": '\n'.join(result_lines)}}
    except Exception as e:
        print(f'[_retrieve_table_data_handler] 错误: {e}', flush=True)
        return {"success": True, "data": {"message": f"表格数据检索失败: {str(e)}"}}


def _find_similar_headers_handler(query_header='', top_k=5, directory_path=''):
    """查找相似表头工具处理"""
    try:
        from nlp.excel.header_embedding import HeaderEmbeddingEncoder
        import os
        
        if not directory_path:
            directory_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
        
        encoder = HeaderEmbeddingEncoder()
        if os.path.exists(directory_path):
            encoder.encode_directory(directory_path)
        
        if not encoder.header_vectors:
            return {"success": True, "data": {"message": "未找到任何表头数据"}}
        
        results = encoder.find_similar_headers(query_header, top_k=top_k)
        
        if not results:
            return {"success": True, "data": {"message": f"未找到与 '{query_header}' 相似的表头"}}
        
        result_lines = [f"与 '{query_header}' 相似的表头（共{len(results)}个）:"]
        for i, res in enumerate(results, 1):
            result_lines.append(f"{i}. {res['header']} (相似度: {res['similarity']:.4f}, 文件: {res['file']})")
        
        return {"success": True, "data": {"message": '\n'.join(result_lines)}}
    except Exception as e:
        print(f'[_find_similar_headers_handler] 错误: {e}', flush=True)
        return {"success": True, "data": {"message": f"查找相似表头失败: {str(e)}"}}


def _calculate_header_similarity_handler(header1='', header2=''):
    """计算表头相似度工具处理"""
    try:
        from nlp.excel.header_embedding import HeaderEmbeddingEncoder
        
        encoder = HeaderEmbeddingEncoder()
        similarity = encoder.calculate_similarity(header1, header2)
        
        if similarity >= 0.8:
            level = "非常相似"
        elif similarity >= 0.6:
            level = "比较相似"
        elif similarity >= 0.4:
            level = "一般相似"
        else:
            level = "不太相似"
        
        return {"success": True, "data": {"message": f"'{header1}' 与 '{header2}' 的相似度: {similarity:.4f} ({level})"}}
    except Exception as e:
        print(f'[_calculate_header_similarity_handler] 错误: {e}', flush=True)
        return {"success": True, "data": {"message": f"计算表头相似度失败: {str(e)}"}}


TOOL_HANDLERS_PERSONAL = {
    'query_knowledge': _query_knowledge_handler,
    'query_organization': _query_organization_handler,
    'web_search': _web_search_handler,
    'generate_image': _generate_image_handler,
    'table_operation': _table_operation_handler,
    'clean_text': _clean_text_handler,
    'extract_text': _extract_text_handler,
    'split_text': _split_text_handler,
    'analyze_table_relationships': _analyze_table_relationships_handler,
    'retrieve_table_data': _retrieve_table_data_handler,
    'find_similar_headers': _find_similar_headers_handler,
    'calculate_header_similarity': _calculate_header_similarity_handler
}


# ========== Format 格式匹配 ==========

FORMAT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'format')
FORMAT_MAP = {
    '文章': 'context.md',
    '论文': 'context.md',
    '小红书': 'red book.md',
    'red book': 'red book.md',
    '公众号': 'Official Account.md',
    'official account': 'Official Account.md',
    '朋友圈': 'friend.md',
    'friend': 'friend.md',
    '视频脚本': 'quick video.md',
    '短视频': 'quick video.md',
    'quick video': 'quick video.md',
    '视频': 'quick video.md'
}

def load_format_file(format_name):
    """加载格式规范文件内容"""
    filename = FORMAT_MAP.get(format_name)
    if not filename:
        return None
    file_path = os.path.join(FORMAT_DIR, filename)
    if not os.path.exists(file_path):
        return None
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f'[load_format_file] 读取失败: {e}')
        return None

def detect_content_type(content):
    """根据内容检测输出类型"""
    content_lower = content.lower()
    type_scores = {
        '小红书': 0,
        '公众号': 0,
        '朋友圈': 0,
        '视频脚本': 0,
        '文章': 0
    }
    
    # 关键词匹配
    if any(k in content_lower for k in ['小红书', 'red book', '种草', '笔记', '博主']):
        type_scores['小红书'] += 10
    if any(k in content_lower for k in ['公众号', '推文', '公众号推文', '微信文章']):
        type_scores['公众号'] += 10
    if any(k in content_lower for k in ['朋友圈', '朋友圈文案', '广告']):
        type_scores['朋友圈'] += 10
    if any(k in content_lower for k in ['视频脚本', '短视频', '分镜头', '脚本', '口播']):
        type_scores['视频脚本'] += 10
    if any(k in content_lower for k in ['文章', '论文', '报告', '总结']):
        type_scores['文章'] += 10
    
    # 返回得分最高的类型
    max_type = max(type_scores, key=type_scores.get)
    if type_scores[max_type] > 0:
        return max_type
    return None


def format_content_with_review(content, format_name, model="qwen-turbo"):
    """
    根据格式规范调整内容，并进行质量评价循环
    返回: (final_content, review_result)
    """
    format_text = load_format_file(format_name)
    if not format_text:
        return content, None
    
    headers = {
        "Authorization": f"Bearer {os.getenv('API_KEY')}",
        "Content-Type": "application/json"
    }
    
    # 第一轮：按格式调整
    is_light = format_name in ['小红书', '朋友圈', '公众号', '视频脚本']
    if is_light:
        adjust_prompt = f"""请根据以下格式规范，对原始内容进行轻量排版优化，使其符合该平台的风格要求。

重要：
1. 必须保留原始内容中的所有实际信息、emoji、短句风格和语气，不要删除或替换为示例文本。
2. 不要添加"根据评价意见修改""严格按照规范调整"等论文式套话。
3. 只进行轻量排版优化（如适当分段、调整标题层级），保持文案的口语化和吸引力。

【格式规范】
{format_text}

【原始内容】
{content}

请直接输出优化后的完整文案，不要添加额外说明。"""
    else:
        adjust_prompt = f"""请根据以下格式规范，将原始内容重新排版和格式化，使其符合规范要求。

重要：必须保留原始内容中的所有实际信息、数据、观点和论述，不要删除、替换为示例文本或生成空框架。你只负责调整排版格式。

【格式规范】
{format_text}

【原始内容】
{content}

请直接输出格式化后的完整内容，不要添加额外说明。"""
    
    try:
        adjust_data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是一位专业的内容编辑，擅长按格式规范调整文本。"},
                {"role": "user", "content": adjust_prompt}
            ],
            "temperature": 0.5,
            "max_tokens": 4096
        }
        response = requests.post("https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions", headers=headers, json=adjust_data, verify=False)
        result = response.json()
        adjusted = result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f'[format_content] 格式调整失败: {e}')
        return content, None
    
    # 质量评价循环（最多3轮）
    max_review_rounds = 3
    current_content = adjusted
    
    for round_num in range(max_review_rounds):
        review_prompt = f"""请对以下内容进行质量评价和完整性评价：

【格式规范】
{format_text}

【待评价内容】
{current_content}

请从以下维度评价：
1. 内容完整性（是否包含必要的章节/要素）
2. 格式符合度（是否符合格式规范要求）
3. 语言质量（表达是否流畅、专业）

如果内容通过评价（所有维度均合格），请回复：「评价通过」
如果内容不通过，请回复：「评价不通过」并列出具体问题和修改建议。"""
        
        try:
            review_data = {
                "model": model,
                "messages": [
                    {"role": "system", "content": "你是一位严格的内容质量评审员。"},
                    {"role": "user", "content": review_prompt}
                ],
                "temperature": 0.3,
                "max_tokens": 2048
            }
            response = requests.post("https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions", headers=headers, json=review_data, verify=False)
            result = response.json()
            review_result = result["choices"][0]["message"]["content"]
            
            if "评价通过" in review_result:
                return current_content, review_result
            
            # 不通过，要求修改
            fix_prompt = f"""请根据以下评价意见修改内容：

【当前内容】
{current_content}

【评价意见】
{review_result}

请输出修改后的完整内容。"""
            
            fix_data = {
                "model": model,
                "messages": [
                    {"role": "system", "content": "你是一位专业的内容编辑，擅长根据反馈修改文本。"},
                    {"role": "user", "content": fix_prompt}
                ],
                "temperature": 0.5,
                "max_tokens": 4096
            }
            response = requests.post("https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions", headers=headers, json=fix_data, verify=False)
            result = response.json()
            current_content = result["choices"][0]["message"]["content"]
            
        except Exception as e:
            print(f'[format_content] 评价第{round_num+1}轮失败: {e}')
            break
    
    return current_content, review_result if 'review_result' in dir() else None


# CORS 配置 - 允许所有来源，但为了支持 credentials，使用自定义函数
def get_cors_origins():
    return [
        'http://localhost:8080',
        'http://127.0.0.1:8080',
        'http://localhost:3315',
        'http://127.0.0.1:3315',
        'http://localhost:5173',
        'http://localhost:3000',
        'http://localhost:5174',
        'http://127.0.0.1:5174',
    ]

CORS(app, resources={
    r"/*": {
        "origins": get_cors_origins(),
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS", "HEAD"],
        "allow_headers": ["*"],
        "supports_credentials": True,
        "expose_headers": ["*"],
        "max_age": 3600
    }
})

register_user_routes(app)

try:
    from ruoyi_langchain.routes.langchain_routes import register_langchain_routes
    register_langchain_routes(app)
    print("[INFO] LangChain路由注册成功")
except ImportError as e:
    print(f"[WARN] LangChain路由注册失败: {e}")
    print("[INFO] 主应用仍可正常运行，LangChain功能将在依赖安装后启用")

try:
    from routes.graph_routes import register_graph_routes
    register_graph_routes(app)
    print("[INFO] GraphRAG路由注册成功")
except ImportError as e:
    print(f"[WARN] GraphRAG路由注册失败: {e}")
    print("[INFO] 主应用仍可正常运行，GraphRAG功能将在依赖安装后启用")

try:
    from routes.multi_agent_routes import register_multi_agent_routes
    register_multi_agent_routes(app)
    print("[INFO] 多Agent协作路由注册成功")
except ImportError as e:
    print(f"[WARN] 多Agent协作路由注册失败: {e}")
    print("[INFO] 主应用仍可正常运行，多Agent功能将在依赖安装后启用")
except Exception as e:
    print(f"[WARN] 多Agent协作路由注册异常: {e}")
    print("[INFO] 主应用仍可正常运行")

@app.route('/')
def hello_world():
    return 'Hello, Flask!'

@app.route('/image/<filename>')
def serve_image(filename):
    from layer.user.upload import serve_image as upload_serve_image
    return upload_serve_image(filename)

@app.route('/media/photo/<filename>')
def serve_media_photo(filename):
    """提供 AI 生成图片的访问路由"""
    from flask import send_from_directory
    import os
    photo_dir = os.path.join(os.path.dirname(__file__), 'old', 'AIuser', 'photo')
    safe_filename = os.path.basename(filename)
    file_path = os.path.join(photo_dir, safe_filename)
    if not os.path.exists(file_path):
        return jsonify({'success': False, 'message': '文件不存在'}), 404
    return send_from_directory(photo_dir, safe_filename)

@app.route('/old/image/<filename>')
def serve_old_image(filename):
    from util.file_utils import read_image_file, is_image_file
    import os
    
    image_dir = os.path.join(os.path.dirname(__file__), 'old', 'image')
    file_path = os.path.join(image_dir, filename)
    
    if not os.path.exists(file_path):
        return '图片不存在', 404
    
    if not is_image_file(filename):
        return '不是有效的图片文件', 400
    
    image_data = read_image_file(file_path)
    
    ext = os.path.splitext(filename)[1].lower()
    mime_types = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.svg': 'image/svg+xml',
        '.webp': 'image/webp'
    }
    
    return image_data, 200, {'Content-Type': mime_types.get(ext, 'image/jpeg')}

@app.route('/api/auth/profile', methods=['GET'])
def get_profile():
    from layer.user.login import get_user_by_account
    user = get_user_by_account('123456')
    
    if user:
        return jsonify({
            'success': True,
            'data': {
                'id': int(user.get('id', 0)),
                'username': user.get('username', ''),
                'account': user.get('account', ''),
                'name': user.get('name', ''),
                'image': user.get('image', ''),
                'created_at': user.get('created_at', '')
            }
        })
    else:
        return jsonify({
            'success': False,
            'message': '用户不存在'
        }), 404

@app.route('/api/agents', methods=['GET'])
def get_agents():
    return jsonify({
        'success': True,
        'data': [
            {
                'id': 'agent-1',
                'name': '豆包',
                'description': '智能助手',
                'status': 'online'
            }
        ]
    })

@app.route('/api/agents/<agent_id>/messages', methods=['POST'])
def send_message(agent_id):
    data = request.get_json()
    message = data.get('message')
    
    return jsonify({
        'success': True,
        'data': {
            'reply': f'您的消息是：{message}\n这是豆包的回复内容。'
        }
    })

@app.route('/api/operations', methods=['GET'])
def get_operations():
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 10, type=int)
    
    return jsonify({
        'success': True,
        'data': {
            'list': [
                {
                    'id': 'REC-2024-001',
                    'name': '用户权限更新',
                    'type': '权限管理',
                    'creator': '张伟',
                    'createTime': '2024-01-15 14:30:22',
                    'status': 'success'
                },
                {
                    'id': 'REC-2024-002',
                    'name': '数据备份任务',
                    'type': '数据管理',
                    'creator': '李娜',
                    'createTime': '2024-01-15 11:20:15',
                    'status': 'processing'
                }
            ],
            'total': 24,
            'page': page,
            'size': size
        }
    })

# 知识库 API 路由
@app.route('/api/knowledge/list', methods=['GET'])
def api_knowledge_list():
    return get_knowledge_list()

@app.route('/api/knowledge/<knowledge_id>', methods=['GET'])
def api_knowledge_detail(knowledge_id):
    return get_knowledge_detail(knowledge_id)

@app.route('/api/knowledge/create', methods=['POST'])
def api_knowledge_create():
    return create_knowledge()

@app.route('/api/knowledge/update/<knowledge_id>', methods=['PUT', 'POST'])
def api_knowledge_update(knowledge_id):
    return update_knowledge(knowledge_id)

@app.route('/api/knowledge/delete/<knowledge_id>', methods=['DELETE', 'POST'])
def api_knowledge_delete(knowledge_id):
    return delete_knowledge(knowledge_id)

@app.route('/api/knowledge/upload', methods=['POST'])
def api_knowledge_upload():
    return upload_attachment()

@app.route('/uploads/knowledge/<filename>')
def api_serve_knowledge_file(filename):
    return serve_attachment(filename)

@app.route('/api/knowledge/statistics', methods=['GET'])
def api_knowledge_statistics():
    return get_statistics()

@app.route('/api/chat', methods=['POST'])
def api_chat():
    data = request.get_json()
    message = data.get('message', '')
    model = data.get('model', 'qwen-turbo')
    conversation_id = data.get('conversation_id')
    history_messages = data.get('history_messages', [])
    file_content = data.get('file_content', '')
    
    import sys
    print(f"[api_chat] 收到请求: message长度={len(message)}, history_messages数量={len(history_messages)}, conversation_id={conversation_id}", flush=True)
    
    if not message:
        return jsonify({
            'success': False,
            'message': '消息不能为空'
        }), 400
    
    try:
        # 如果请求包含文件内容（上传文件），先解析并存入会话，供后续@个人助手使用
        if file_content and conversation_id:
            _parsed = _parse_file_content(file_content)
            if _parsed:
                session_table_data[conversation_id] = _parsed
                _hdrs = _parsed.get('headers')
                print(f'[api_chat] 文件内容已存入session: conv={conversation_id}, headers={_hdrs}', flush=True)
        
        # 检测 @前缀，选择对应的工具和调用方式
        if '@知识库' in message:
            response, conv_id = call_api_with_tools(
                message, TOOLS_KNOWLEDGE, TOOL_HANDLERS_KNOWLEDGE, model=model,
                save=True, conversation_id=conversation_id,
                history_messages=history_messages, display_prompt=message
            )
            return jsonify({
                'success': True,
                'data': {
                    'reply': response,
                    'conversation_id': conv_id
                }
            })
        
        elif '@组织管理' in message or '@组织架构' in message:
            response, conv_id = call_api_with_tools(
                message, TOOLS_ORGANIZATION, TOOL_HANDLERS_ORGANIZATION, model=model,
                save=True, conversation_id=conversation_id,
                history_messages=history_messages, display_prompt=message
            )
            return jsonify({
                'success': True,
                'data': {
                    'reply': response,
                    'conversation_id': conv_id
                }
            })
        
        elif '@联网搜索' in message:
            # 提取搜索关键词（去掉 @联网搜索 前缀）
            query = message.replace('@联网搜索', '').strip()
            if not query:
                return jsonify({
                    'success': False,
                    'message': '请输入搜索关键词'
                }), 400
            
            # 直接调用博查搜索 API，返回真实搜索结果
            search_result = call_search_api(query)
            
            if search_result.get('success'):
                data = search_result.get('data', {})
                web_pages = data.get('webPages', {})
                results = web_pages.get('value', [])
                
                # 格式化搜索结果为 Markdown
                lines = [f'🔍 **「{query}」的博查搜索结果**', '']
                if results:
                    for idx, item in enumerate(results[:10], 1):
                        title = item.get('name', '无标题')
                        url = item.get('url', '')
                        snippet = item.get('snippet', '')
                        lines.append(f'{idx}. **{title}**')
                        if snippet:
                            lines.append(f'   {snippet}')
                        if url:
                            lines.append(f'   [{url}]({url})')
                        lines.append('')
                else:
                    lines.append('未找到相关结果。')
                
                response = '\n'.join(lines)
            else:
                response = f'搜索失败：{search_result.get("message", "未知错误")}'
            
            # 生成 conversation_id 用于保存历史
            conv_id = conversation_id if conversation_id else datetime.now().strftime('%Y%m%d_%H%M%S')
            
            return jsonify({
                'success': True,
                'data': {
                    'reply': response,
                    'conversation_id': conv_id
                }
            })
        
        elif '@个人助手' in message:
            print(f'[个人助手] 进入分支, file_content长度: {len(file_content) if file_content else 0}', flush=True)
            # 从消息中提取用户原始问题（去掉文件内容部分）
            user_query = message
            if '---\n用户问题：' in message:
                user_query = message.split('---\n用户问题：')[-1].strip()
            elif '---\r\n用户问题：' in message:
                user_query = message.split('---\r\n用户问题：')[-1].strip()
            
            # 如果消息中包含本地文件路径且没有file_content，尝试自动读取
            if not file_content:
                import re as _re_path
                _q = user_query.replace('@个人助手', '')
                # 路径匹配：匹配 Windows 路径（盘符开头，到空格/换行为止）
                path_match = _re_path.search(
                    r'([A-Za-z]:[\\\/][^\s\n]+)',
                    _q
                )
                if path_match:
                    raw_path = path_match.group(1).rstrip('。，！？；：""''（）【】、')
                    # 尝试找到存在的文件路径
                    found_path = None
                    
                    # 策略：按路径分隔符逐步截断，对最后一段尝试去掉中文描述后缀
                    segments = raw_path.replace('/', '\\').split('\\')
                    for i in range(len(segments), 1, -1):
                        test_path = '\\'.join(segments[:i])
                        # 尝试直接路径
                        if os.path.isfile(test_path):
                            found_path = test_path
                            break
                        # 尝试加扩展名
                        base, ext = os.path.splitext(test_path)
                        if not ext:
                            for _ext in ['.xlsx', '.xls', '.csv', '.txt', '.json']:
                                if os.path.isfile(test_path + _ext):
                                    found_path = test_path + _ext
                                    break
                        if found_path:
                            break
                        
                        # 对最后一段，尝试去掉中文描述后缀（从右往左逐字去掉）
                        last_seg = segments[i-1]
                        parent = '\\'.join(segments[:i-1]) if i > 1 else ''
                        if parent and len(last_seg) > 1:
                            # 逐步缩短最后一段，尝试匹配文件
                            for cut_pos in range(len(last_seg)-1, 0, -1):
                                short_seg = last_seg[:cut_pos]
                                test_name = os.path.join(parent, short_seg)
                                if os.path.isfile(test_name):
                                    found_path = test_name
                                    break
                                base2, ext2 = os.path.splitext(test_name)
                                if not ext2:
                                    for _ext in ['.xlsx', '.xls', '.csv', '.txt', '.json']:
                                        if os.path.isfile(test_name + _ext):
                                            found_path = test_name + _ext
                                            break
                                if found_path:
                                    break
                        if found_path:
                            break
                    
                    _log3 = f'[个人助手] 路径匹配: raw_path={raw_path!r}, found_path={found_path!r}'
                    print(_log3, flush=True)
                    
                    if found_path:
                        try:
                            _ext = os.path.splitext(found_path)[1].lower()
                            if _ext in ['.xlsx', '.xls']:
                                import pandas as _pd_read
                                _df = _pd_read.read_excel(found_path, engine='openpyxl' if _ext == '.xlsx' else 'xlrd')
                                file_content = _df.to_csv(index=False, sep='\t')
                            elif _ext == '.csv':
                                import pandas as _pd_read
                                _df = _pd_read.read_csv(found_path, encoding='utf-8-sig')
                                file_content = _df.to_csv(index=False, sep='\t')
                            elif _ext == '.txt':
                                with open(found_path, 'r', encoding='utf-8-sig') as _f:
                                    file_content = _f.read()
                            elif _ext == '.json':
                                with open(found_path, 'r', encoding='utf-8-sig') as _f:
                                    file_content = _f.read()
                            print(f'[个人助手] 自动读取文件: {found_path}，长度: {len(file_content)}')
                        except Exception as _e:
                            print(f'[个人助手] 读取文件失败 {found_path}: {_e}')

            # 检查是否为表格操作命令
            table_action = _parse_table_action(user_query.replace('@个人助手', '').strip())
            if table_action:
                action, params = table_action
                
                # 1. 优先使用请求中的 file_content 字段
                parsed_table_data = None
                if file_content:
                    parsed_table_data = _parse_file_content(file_content)
                    print(f'[个人助手] file_content 解析结果: headers={parsed_table_data.get("headers") if parsed_table_data else None}')
                
                # 2. 若 file_content 为空，尝试从消息文本中的代码块提取表格内容
                if not parsed_table_data and message:
                    import re as _re_block
                    block_match = _re_block.search(r'```[\s\S]*?\n([\s\S]+?)\n```', message)
                    if block_match:
                        block_content = block_match.group(1)
                        parsed_table_data = _parse_file_content(block_content)
                        print(f'[个人助手] 消息代码块解析结果: headers={parsed_table_data.get("headers") if parsed_table_data else None}')
                
                # 3. 存入 session，并使用刚解析的数据（避免 session 丢失导致 None）
                if parsed_table_data:
                    if conversation_id:
                        session_table_data[conversation_id] = parsed_table_data
                    table_data = parsed_table_data
                else:
                    # 回退：从 session 中取历史表格数据
                    table_data = session_table_data.get(conversation_id) if conversation_id else None
                
                if not table_data:
                    print(f'[个人助手] 警告：table_data 为空，无法执行图表操作')
                
                # 执行表格操作
                result = call_table_tool(action, table_data=table_data, **params)
                
                # 格式化结果
                reply = _format_table_result(result)
                # 如果包含HTML图表，标记is_html
                is_html_reply = '<div class="chart-wrapper"' in reply or '<img ' in reply
                
                # 更新会话表格数据
                if result.get('table_data') and conversation_id:
                    session_table_data[conversation_id] = result['table_data']
                
                conv_id = conversation_id if conversation_id else datetime.now().strftime('%Y%m%d_%H%M%S')
                
                reply_payload = {'reply': reply, 'conversation_id': conv_id}
                if is_html_reply:
                    reply_payload['is_html'] = True
                
                return jsonify({
                    'success': True,
                    'data': reply_payload
                })
            
            # 非表格操作，使用 LLM 工具调用
            # 确保 conversation_id 存在，用于 session 表格数据存储
            if not conversation_id:
                conversation_id = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # 从历史消息中提取之前生成的图片URL和图表信息
            history_images = []
            import re as _re_img
            for hmsg in history_messages:
                if hmsg.get('role') == 'assistant':
                    content = hmsg.get('content', '')
                    # 提取 markdown 图片链接 ![alt](url)
                    md_imgs = _re_img.findall(r'!\[([^\]]*)\]\(([^)]+)\)', content)
                    for alt, url in md_imgs:
                        history_images.append({'alt': alt, 'url': url})
                    # 提取 HTML img 标签
                    html_imgs = _re_img.findall(r'<img[^>]+src=["\']([^"\']+)["\']', content)
                    for url in html_imgs:
                        history_images.append({'alt': '图表/图片', 'url': url})
            
            # 构建历史图片上下文
            image_context = ''
            if history_images:
                image_context = '\n\n【之前对话中生成的图片和图表】\n'
                for idx, img in enumerate(history_images, 1):
                    image_context += f'{idx}. {img["alt"]}: ![{img["alt"]}]({img["url"]})\n'
                image_context += '\n请在生成文章或文案时，引用上述图片（使用相同的Markdown格式插入）。\n'
            
            # 如果有 file_content，解析表格数据存入 session，并注入到消息中供 LLM 分析
            if file_content:
                parsed_table = _parse_file_content(file_content)
                if parsed_table and conversation_id:
                    session_table_data[conversation_id] = parsed_table
                    print(f'[个人助手] 文件内容已解析为表格数据并存入session, headers={parsed_table.get("headers")}')
                # 发送完整文件内容，不截断（避免分析只看到部分公司/数据）
                llm_message = f"{message}\n\n【文件数据内容（完整）】\n{file_content}{image_context}"
                print(f'[个人助手] 注入文件内容到LLM消息, 总长度: {len(file_content)}, 历史图片数: {len(history_images)}')
            else:
                llm_message = f"{message}{image_context}"
                if history_images:
                    print(f'[个人助手] 注入历史图片上下文, 图片数: {len(history_images)}')
            
            response, conv_id = call_api_with_tools(
                llm_message, TOOLS_PERSONAL_ASSISTANT, TOOL_HANDLERS_PERSONAL, model=model,
                save=True, conversation_id=conversation_id,
                history_messages=history_messages, display_prompt=message
            )
            
            # 兜底：检测AI是否编造了不存在的图片路径，若编造则自动调用 generate_image 生成真实图片
            import re as _img_re
            fake_img_pattern = r'!\[([^\]]*)\]\(/media/photo/([a-f0-9]{32}\.png)\)'
            fake_matches = _img_re.findall(fake_img_pattern, response)
            if fake_matches:
                print(f'[个人助手] 检测到 {len(fake_matches)} 个编造图片路径，自动调用 generate_image 生成')
                for alt_text, fake_filename in fake_matches:
                    # 检查文件是否真实存在（图片保存在 old/AIuser/photo/）
                    fake_path = os.path.join(os.path.dirname(__file__), 'old', 'AIuser', 'photo', fake_filename)
                    if not os.path.exists(fake_path):
                        # 自动调用图片生成（使用 handler 轮询等待完成）
                        gen_result = _generate_image_handler(prompt=alt_text or '相关图片')
                        if gen_result.get('success') and gen_result.get('data', {}).get('images'):
                            real_url = gen_result['data']['images'][0].get('local_url', '')
                            if real_url:
                                response = response.replace(
                                    f'![{alt_text}](/media/photo/{fake_filename})',
                                    f'![{alt_text}]({real_url})'
                                )
                                print(f'[个人助手] 已替换编造图片: {fake_filename} -> {real_url}')
                        else:
                            _err_msg = gen_result.get('message', '未知错误')
                            print(f'[个人助手] 图片生成失败: {_err_msg}')
            
            # 检测内容类型并进行 Format 匹配和质量评价
            # 优先根据用户原始消息判断类型，避免AI回复中缺少关键词导致误判
            user_msg_lower = message.lower()
            if '小红书' in user_msg_lower or 'red book' in user_msg_lower:
                content_type = '小红书'
            elif '公众号' in user_msg_lower or '推文' in user_msg_lower:
                content_type = '公众号'
            elif '朋友圈' in user_msg_lower:
                content_type = '朋友圈'
            elif '视频脚本' in user_msg_lower or '短视频' in user_msg_lower or '脚本' in user_msg_lower:
                content_type = '视频脚本'
            else:
                content_type = detect_content_type(response)
            if content_type:
                response, review = format_content_with_review(response, content_type, model)
            
            # 检测回复中是否包含图片URL，如果有则提取为附件并转换为HTML显示
            import re
            image_pattern = r'https?://[^\s\"\'<>]+\.(?:png|jpg|jpeg|gif|webp)(?:\?[^\s\"\'<>]*)?'
            image_urls = re.findall(image_pattern, response, re.IGNORECASE)
            
            reply_data = {
                'reply': response,
                'conversation_id': conv_id
            }
            
            if image_urls:
                # 下载图片到本地保存
                img_dir = os.path.join(os.path.dirname(__file__), 'uploads', 'chat_images')
                os.makedirs(img_dir, exist_ok=True)
                local_attachments = []
                url_mapping = {}  # 外部URL -> 本地URL 映射
                
                for url in image_urls:
                    try:
                        img_resp = requests.get(url, timeout=60, verify=False)
                        if img_resp.status_code == 200:
                            ext = os.path.splitext(url.split('?')[0])[1] or '.png'
                            if ext not in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
                                ext = '.png'
                            local_filename = f"{uuid.uuid4().hex}{ext}"
                            local_path = os.path.join(img_dir, local_filename)
                            with open(local_path, 'wb') as f:
                                f.write(img_resp.content)
                            local_url = f'/uploads/chat_images/{local_filename}'
                            url_mapping[url] = local_url
                            local_attachments.append({'type': 'image', 'url': local_url})
                    except Exception as e:
                        print(f'[WARN] 下载图片失败: {e}')
                        local_attachments.append({'type': 'image', 'url': url})
                
                # 将 response 中的外部URL替换为本地URL
                for ext_url, local_url in url_mapping.items():
                    response = response.replace(ext_url, local_url)
                
                # 将Markdown链接转换为直接图片显示（使用本地URL）
                md_img_pattern = r'\!\[([^\]]*)\]\((https?://[^\s\"\'<>]+\.(?:png|jpg|jpeg|gif|webp)(?:\?[^\s\"\'<>]*)?)\)'
                response = re.sub(md_img_pattern, r'<img src="\2" alt="\1" class="max-w-xs max-h-48 rounded-lg border mb-2 cursor-zoom-in" onclick="openImageLightbox(\'\2\')">', response, flags=re.IGNORECASE)
                
                md_link_pattern = r'\[([^\]]+)\]\((https?://[^\s\"\'<>]+\.(?:png|jpg|jpeg|gif|webp)(?:\?[^\s\"\'<>]*)?)\)'
                response = re.sub(md_link_pattern, r'<img src="\2" alt="\1" class="max-w-xs max-h-48 rounded-lg border mb-2 cursor-zoom-in" onclick="openImageLightbox(\'\2\')">', response, flags=re.IGNORECASE)
                
                bare_url_pattern = r'(?<![\"\'=\(])https?://[^\s\"\'<>]+\.(?:png|jpg|jpeg|gif|webp)(?:\?[^\s\"\'<>]*)?'
                response = re.sub(bare_url_pattern, lambda m: f'<img src="{m.group(0)}" alt="生成的图片" class="max-w-xs max-h-48 rounded-lg border mb-2 cursor-zoom-in" onclick="openImageLightbox(\'{m.group(0)}\')">', response, flags=re.IGNORECASE)
                
                reply_data['reply'] = response
                reply_data['is_html'] = True
                reply_data['attachments'] = local_attachments if local_attachments else [{'type': 'image', 'url': url} for url in image_urls]
            
            return jsonify({
                'success': True,
                'data': reply_data
            })
        
        else:
            # 普通对话，不使用 Function Calling
            print(f'[api_chat] 普通对话分支, message前20字符: {message[:20]!r}', flush=True)
            response, conv_id = call_api(
                message, model, save=True, conversation_id=conversation_id,
                history_messages=history_messages
            )
            return jsonify({
                'success': True,
                'data': {
                    'reply': response,
                    'conversation_id': conv_id
                }
            })
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/chat/image', methods=['POST'])
def api_chat_image():
    data = request.get_json()
    image_url = data.get('image_url', '')
    original_question = data.get('question', '这张图片是什么？')
    question = enhance_prompt(original_question)
    model = data.get('model', 'qwen-vl-plus')
    conversation_id = data.get('conversation_id')
    history_messages = data.get('history_messages', [])
    
    if not image_url:
        return jsonify({
            'success': False,
            'message': '图片URL不能为空'
        }), 400
    
    try:
        import requests
        import base64
        import uuid
        from dotenv import load_dotenv
        load_dotenv('old/.env')
        API_KEY = os.getenv('API_KEY')
        
        image_save_path = ''
        final_image_url = image_url
        
        if image_url.startswith('data:image/'):
            img_data = image_url.split(',')[1]
            img_format = image_url.split(';')[0].split('/')[1]
            img_bytes = base64.b64decode(img_data)
            
            save_dir = os.path.join(os.path.dirname(__file__), 'old', 'image')
            os.makedirs(save_dir, exist_ok=True)
            
            filename = f"{uuid.uuid4().hex}.{img_format}"
            image_save_path = os.path.join(save_dir, filename)
            
            with open(image_save_path, 'wb') as f:
                f.write(img_bytes)
        elif image_url.startswith('http://') or image_url.startswith('https://'):
            # 如果是本地服务器URL，尝试读取本地文件并转为base64
            try:
                local_path = None
                if '/uploads/chat_images/' in image_url:
                    parts = image_url.split('/uploads/chat_images/')
                    if len(parts) == 2:
                        local_path = os.path.join(os.path.dirname(__file__), 'uploads', 'chat_images', parts[1])
                if local_path and os.path.exists(local_path):
                    with open(local_path, 'rb') as f:
                        img_bytes = f.read()
                    ext = os.path.splitext(local_path)[1].lower()
                    mime = 'image/png' if ext == '.png' else 'image/jpeg'
                    b64 = base64.b64encode(img_bytes).decode('utf-8')
                    final_image_url = f"data:{mime};base64,{b64}"
                    image_save_path = local_path
            except Exception as e:
                print(f'[WARN] 转换本地图片为base64失败: {e}')
        
        headers = {
            'Authorization': f'Bearer {API_KEY}',
            'Content-Type': 'application/json'
        }
        
        # 构建消息列表，包含历史对话
        messages = []
        if history_messages:
            for msg in history_messages:
                clean_msg = {"role": msg["role"], "content": msg["content"]}
                messages.append(clean_msg)
            print(f"[api_chat_image] history_messages 数量: {len(history_messages)}")
        else:
            print(f"[api_chat_image] history_messages 为空")
        
        # 当前用户消息包含图片
        messages.append({
            'role': 'user',
            'content': [
                {'type': 'text', 'text': question},
                {'type': 'image_url', 'image_url': {'url': final_image_url}}
            ]
        })
        
        req_data = {
            'model': model,
            'messages': messages
        }
        
        response = requests.post(
            'https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions',
            headers=headers,
            json=req_data,
            verify=False
        )
        result = response.json()
        
        if 'choices' in result:
            content = result['choices'][0]['message']['content']
            conv_id = conversation_id if conversation_id else datetime.now().strftime('%Y%m%d_%H%M%S')
            return jsonify({
                'success': True,
                'data': {
                    'reply': content,
                    'conversation_id': conv_id,
                    'saved_path': image_save_path if image_save_path else None
                }
            })
        else:
            return jsonify({
                'success': False,
                'message': result.get('error', {}).get('message', '未知错误')
            }), 500
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/file/parse', methods=['POST'])
def api_file_parse():
    """解析上传的数据文件（Excel/CSV/TXT/JSON），返回文本内容"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '没有上传文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '文件名为空'}), 400
    
    try:
        filename = file.filename
        ext = os.path.splitext(filename)[1].lower()
        
        if ext in ['.csv']:
            import csv
            import io
            content = []
            stream = io.StringIO(file.stream.read().decode('utf-8-sig'), newline=None)
            reader = csv.reader(stream)
            for row in reader:
                content.append('\t'.join(row))
            text = '\n'.join(content)
            
        elif ext in ['.xlsx', '.xls']:
            import pandas as pd
            df = pd.read_excel(file, engine='openpyxl' if ext == '.xlsx' else 'xlrd')
            # 直接返回表格数据结构，避免格式转换问题
            headers = [str(col) for col in df.columns]
            rows = []
            for _, row in df.iterrows():
                rows.append([str(cell) for cell in row.values])
            return jsonify({
                'success': True,
                'data': {
                    'content': '',
                    'table_data': {
                        'headers': headers,
                        'rows': rows,
                        'row_count': len(rows),
                        'col_count': len(headers),
                        'table_name': filename
                    }
                }
            })
            
        elif ext in ['.json']:
            import json
            data = json.load(file)
            text = json.dumps(data, ensure_ascii=False, indent=2)
            
        elif ext in ['.txt', '.md', '.py']:
            text = file.read().decode('utf-8')
            
        else:
            return jsonify({'success': False, 'message': f'不支持的文件格式: {ext}'}), 400
        
        max_len = 30000
        if len(text) > max_len:
            text = text[:max_len] + f'\n\n...（内容已截断，原始大小 {len(text)} 字符）'
        
        return jsonify({
            'success': True,
            'data': {
                'filename': filename,
                'content': text,
                'length': len(text)
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'解析失败: {str(e)}'}), 500


@app.route('/uploads/chat_images/<filename>')
def serve_chat_image(filename):
    """提供聊天图片的本地访问"""
    img_dir = os.path.join(os.path.dirname(__file__), 'uploads', 'chat_images')
    return send_from_directory(img_dir, filename)


@app.route('/api/upload/latest', methods=['GET'])
def api_upload_latest():
    """获取最新扫码上传的文件内容"""
    try:
        files = os.listdir(UPLOAD_TEMP_DIR)
        if not files:
            return jsonify({'success': False, 'message': '没有上传的文件'}), 404
        
        files.sort(key=lambda f: os.path.getmtime(os.path.join(UPLOAD_TEMP_DIR, f)), reverse=True)
        latest_file = files[0]
        filepath = os.path.join(UPLOAD_TEMP_DIR, latest_file)
        
        filename = latest_file
        if '_' in filename and filename.split('_')[0].isdigit():
            filename = '_'.join(filename.split('_')[1:])
        
        ext = os.path.splitext(filename)[1].lower()
        
        if ext in ['.csv']:
            import csv
            content = []
            with open(filepath, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                for row in reader:
                    content.append(', '.join(row))
            text = '\n'.join(content)
            file_type = 'csv'
        elif ext in ['.xlsx', '.xls']:
            import pandas as pd
            df = pd.read_excel(filepath, engine='openpyxl' if ext == '.xlsx' else 'xlrd')
            text = df.to_string(index=False)
            file_type = 'excel'
        elif ext in ['.json']:
            import json
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            text = json.dumps(data, ensure_ascii=False, indent=2)
            file_type = 'json'
        elif ext in ['.txt', '.md']:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
            file_type = 'text'
        elif ext in ['.doc', '.docx']:
            try:
                import docx
                doc = docx.Document(filepath)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = '\n'.join(paragraphs)
                file_type = 'doc'
            except ImportError:
                return jsonify({'success': False, 'message': '缺少 python-docx 库，无法读取Word文档'}), 500
        elif ext in ['.jpg', '.jpeg', '.png']:
            # 图片文件：复制到可访问目录并返回URL
            img_dir = os.path.join(os.path.dirname(__file__), 'uploads', 'chat_images')
            os.makedirs(img_dir, exist_ok=True)
            import uuid
            img_ext = ext
            new_filename = f"qr_{uuid.uuid4().hex}{img_ext}"
            new_path = os.path.join(img_dir, new_filename)
            import shutil
            shutil.copy2(filepath, new_path)
            return jsonify({
                'success': True,
                'data': {
                    'filename': filename,
                    'type': 'image',
                    'url': f'/uploads/chat_images/{new_filename}'
                }
            })
        else:
            return jsonify({'success': False, 'message': f'不支持的文件格式: {ext}'}), 400
        
        max_len = 30000
        if len(text) > max_len:
            text = text[:max_len] + f'\n\n...（内容已截断，原始大小 {len(text)} 字符）'
        
        return jsonify({
            'success': True,
            'data': {
                'filename': filename,
                'content': text,
                'type': file_type,
                'length': len(text)
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/chat/history', methods=['GET'])
def api_chat_history():
    import os
    import json as _json
    
    data_dir = os.path.join(os.path.dirname(__file__), 'old', 'data')
    
    if not os.path.exists(data_dir):
        return jsonify({
            'success': True,
            'data': []
        })
    
    history = []
    # 同时读取 json 和 txt，json 优先
    all_files = []
    for f in os.listdir(data_dir):
        if f.endswith('.json'):
            all_files.append((f, 'json'))
        elif f.endswith('.txt'):
            # 如果有同名的 json，跳过 txt
            json_name = f.replace('.txt', '.json')
            if not os.path.exists(os.path.join(data_dir, json_name)):
                all_files.append((f, 'txt'))
    
    # 按文件名倒序（最新的在前面）
    all_files.sort(key=lambda x: x[0], reverse=True)
    
    for filename, ftype in all_files[:50]:
        file_path = os.path.join(data_dir, filename)
        try:
            if ftype == 'json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = _json.load(f)
                
                qa_pairs = data.get('qa_pairs', [])
                if qa_pairs:
                    history.append({
                        'id': filename.replace('.json', ''),
                        'qa_pairs': qa_pairs,
                        'first_question': qa_pairs[0].get('question', ''),
                        'last_answer': qa_pairs[-1].get('answer', '')[:200],
                        'time': data.get('time', ''),
                        'turns': len(qa_pairs)
                    })
            else:
                # txt 格式回退兼容
                with open(file_path, 'r', encoding='utf-8-sig') as f:
                    content = f.read()
                
                lines = content.strip().split('\n')
                time_str = ''
                qa_pairs = []
                current_question = ''
                current_answer_lines = []
                
                for line in lines:
                    if line.startswith('时间:'):
                        time_str = line[3:].strip()
                    elif line.startswith('用户:'):
                        if current_question and current_answer_lines:
                            qa_pairs.append({
                                'question': current_question,
                                'answer': '\n'.join(current_answer_lines),
                                'question_attachments': [],
                                'answer_attachments': []
                            })
                        current_question = line[3:].strip()
                        current_answer_lines = []
                    elif line.startswith('RYagent:'):
                        current_answer_lines.append(line[8:].strip())
                    elif current_question and current_answer_lines and line and not line.startswith('=') and not line.startswith('对话ID'):
                        current_answer_lines.append(line)
                
                if current_question and current_answer_lines:
                    qa_pairs.append({
                        'question': current_question,
                        'answer': '\n'.join(current_answer_lines),
                        'question_attachments': [],
                        'answer_attachments': []
                    })
                
                if qa_pairs:
                    history.append({
                        'id': filename.replace('.txt', ''),
                        'qa_pairs': qa_pairs,
                        'first_question': qa_pairs[0]['question'],
                        'last_answer': qa_pairs[-1]['answer'][:200],
                        'time': time_str,
                        'turns': len(qa_pairs)
                    })
        except Exception as e:
            continue
    
    return jsonify({
        'success': True,
        'data': history
    })

@app.route('/api/chat/history', methods=['POST'])
def api_save_history():
    """保存历史记录"""
    import os
    import json as _json

    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': '无效的数据'}), 400

    history_id = data.get('id')
    qa_pairs = data.get('qa_pairs', [])
    if not history_id or not qa_pairs:
        return jsonify({'success': False, 'message': '缺少必要字段'}), 400

    data_dir = os.path.join(os.path.dirname(__file__), 'old', 'data')
    os.makedirs(data_dir, exist_ok=True)

    file_path = os.path.join(data_dir, f'{history_id}.json')
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            _json.dump({
                'time': data.get('time', ''),
                'conversation_id': history_id,
                'qa_pairs': qa_pairs
            }, f, ensure_ascii=False, indent=2)
        return jsonify({'success': True, 'message': '保存成功'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/chat/history/<history_id>', methods=['DELETE'])
def api_delete_history(history_id):
    import os
    
    data_dir = os.path.join(os.path.dirname(__file__), 'old', 'data')
    
    # 优先删除 json，如果没有则删除 txt
    json_path = os.path.join(data_dir, f'{history_id}.json')
    txt_path = os.path.join(data_dir, f'{history_id}.txt')
    
    file_path = None
    if os.path.exists(json_path):
        file_path = json_path
    elif os.path.exists(txt_path):
        file_path = txt_path
    
    if not file_path:
        return jsonify({
            'success': False,
            'message': '记录不存在'
        }), 404
    
    try:
        os.remove(file_path)
        # 如果删除的是 json，同时尝试删除同名的 txt
        if file_path == json_path and os.path.exists(txt_path):
            os.remove(txt_path)
        return jsonify({
            'success': True,
            'message': '删除成功'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

# ==================== 手机连接相关接口 ====================

@app.route('/api/qrcode/server-url', methods=['GET'])
def api_qrcode_server_url():
    """获取服务器局域网IP和连接状态"""
    try:
        # 使用缓存的局域网IP
        local_ip = get_local_ip()
        
        # 清理超时连接（心跳检测），10秒无心跳即断开
        check_expired_users(timeout_seconds=10)
        
        # 检查是否有已连接的用户
        users = get_connected_users()
        if users:
            user = users[0]
            return jsonify({
                'success': True,
                'connected': True,
                'username': user['user_name'],
                'user_id': user['user_id'],
                'server_url': f'http://{local_ip}:5000'
            })
        
        return jsonify({
            'success': True,
            'connected': False,
            'server_url': f'http://{local_ip}:5000'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/qrcode/heartbeat', methods=['POST'])
def api_qrcode_heartbeat():
    """接收客户端心跳"""
    data = request.get_json() or {}
    user_id = data.get('user_id')
    
    if not user_id:
        return jsonify({
            'success': False,
            'message': '缺少user_id'
        }), 400
    
    result = update_heartbeat(user_id)
    return jsonify(result)

@app.route('/api/phone/send', methods=['POST'])
def api_phone_send():
    """发送内容给已连接的手机"""
    data = request.get_json() or {}
    users = get_connected_users()
    
    if not users:
        return jsonify({
            'success': False,
            'message': '没有已连接的手机'
        }), 400
    
    user_id = users[0]['user_id']
    result = send_to_connected_user(user_id, data, data.get('format', 'formatted'))
    return jsonify(result)

@app.route('/api/phone/messages', methods=['GET'])
def api_phone_messages():
    """获取发送给手机的消息"""
    user_id = request.args.get('user_id')
    
    if not user_id:
        return jsonify({
            'success': False,
            'message': '缺少user_id'
        }), 400
    
    messages = get_unread_messages(user_id)
    return jsonify({
        'success': True,
        'has_new': len(messages) > 0,
        'messages': messages
    })

@app.route('/api/qrcode/disconnect', methods=['POST'])
def api_qrcode_disconnect():
    """断开手机连接"""
    data = request.get_json() or {}
    user_id = data.get('user_id')
    
    if user_id:
        result = disconnect_user(user_id)
    else:
        # 断开所有连接
        users = get_connected_users()
        for user in users:
            disconnect_user(user['user_id'])
        result = {'success': True, 'message': '已断开所有连接'}
    
    return jsonify(result)

@app.route('/mobile/connect', methods=['GET', 'POST'])
def mobile_connect():
    """手机扫码连接页面"""
    if request.method == 'POST':
        data = request.get_json() or request.form.to_dict()
        user_name = data.get('name', '未知用户')
        user_id = data.get('user_id', 'mobile_' + str(int(__import__('time').time())))
        result = register_connected_user(user_id, user_name)
        return jsonify(result)
    
    # GET 请求返回HTML页面
    return '''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate">
    <title>手机连接</title>
    <style>
        * { box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; width: 100%; min-height: 100vh; margin: 0; padding: 20px; background: #fff; color: #000; }
        .card { background: #f9f9f9; border: 1px solid #e5e5e5; border-radius: 12px; padding: 24px; margin-bottom: 16px; }
        h2 { text-align: center; color: #000; margin-bottom: 24px; }
        h3 { margin: 0; font-size: 16px; color: #000; }
        input { width: 100%; padding: 12px; margin: 8px 0; border: 1px solid #ccc; border-radius: 8px; box-sizing: border-box; font-size: 16px; background: #fff; color: #000; }
        button { width: 100%; padding: 14px; margin-top: 16px; border: 1px solid #000; border-radius: 8px; font-size: 16px; cursor: pointer; }
        .btn-primary { background: #fff; color: #000; }
        .btn-primary:active { background: #eee; }
        .btn-danger { background: #fff; color: #000; }
        .btn-danger:active { background: #eee; }
        .btn-download { background: #fff; color: #000; padding: 8px 14px; font-size: 14px; margin-top: 8px; border: 1px solid #000; }
        .status { text-align: center; color: #666; margin-bottom: 16px; }
        #result { margin-top: 16px; padding: 12px; border-radius: 8px; text-align: center; display: none; }
        .success { background: #f0f0f0; color: #000; border: 1px solid #ccc; }
        .error { background: #f0f0f0; color: #cc0000; border: 1px solid #ccc; }
        .info-box { background: #f0f0f0; border: 1px solid #ccc; border-radius: 8px; padding: 12px; margin-top: 12px; text-align: center; color: #000; }
        .disconnected { background: #f0f0f0; color: #999; }
        .message-card { background: #f9f9f9; border: 1px solid #e5e5e5; border-radius: 12px; padding: 16px; margin-bottom: 12px; }
        .message-title { font-size: 14px; font-weight: bold; color: #000; margin-bottom: 8px; }
        .message-text { font-size: 14px; color: #000; line-height: 1.6; white-space: pre-wrap; word-break: break-all; background: #f0f0f0; padding: 10px; border-radius: 8px; margin-bottom: 8px; }
        .message-image { max-width: 100%; border-radius: 8px; margin-bottom: 8px; }
        .message-time { font-size: 12px; color: #999; text-align: right; }
        .badge { display: inline-block; background: #000; color: #fff; font-size: 12px; padding: 2px 8px; border-radius: 10px; margin-left: 8px; }
    </style>
</head>
<body>
    <div class="card">
        <h2>连接确认</h2>
        <p class="status" id="status">请输入您的信息</p>
        <div id="inputForm">
            <input type="text" id="userName" placeholder="您的姓名" value="手机用户">
            <button id="connectBtn" class="btn-primary" onclick="connect()">确认连接</button>
        </div>
        <div id="connectedView" style="display:none;">
            <div class="info-box">
                <div id="connStatus">已连接到电脑</div>
                <div style="font-size: 12px; color: #ccc; margin-top: 4px;">保持此页面打开以维持连接</div>
            </div>
            <button id="disconnectBtn" class="btn-danger" onclick="disconnect()" style="margin-top: 16px;">断开连接</button>
        </div>
        <div id="result"></div>
    </div>
    <div id="messageArea" style="display:none;">
        <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:12px;">
            <h3>收到的内容</h3>
            <span id="newBadge" class="badge" style="display:none;">新消息</span>
        </div>
        <div id="messageList"></div>
    </div>
    <script>
        let currentUserId = null;
        let heartbeatTimer = null;
        let messageTimer = null;
        let hasNewMessage = false;
        
        async function connect() {
            const name = document.getElementById("userName").value.trim();
            if (!name) { alert("请输入姓名"); return; }
            
            const btn = document.getElementById("connectBtn");
            btn.textContent = "连接中...";
            btn.disabled = true;
            
            try {
                const controller = new AbortController();
                setTimeout(() => controller.abort(), 8000);
                
                const response = await fetch("/mobile/connect", {
                    method: "POST",
                    headers: { "Content-Type": "application/json" },
                    body: JSON.stringify({ name: name }),
                    signal: controller.signal
                });
                const result = await response.json();
                
                const resultDiv = document.getElementById("result");
                resultDiv.style.display = "block";
                if (result.success) {
                    resultDiv.className = "success";
                    resultDiv.textContent = "连接成功！";
                    document.getElementById("inputForm").style.display = "none";
                    document.getElementById("connectedView").style.display = "block";
                    document.getElementById("messageArea").style.display = "block";
                    document.getElementById("status").textContent = "已连接到服务器";
                    currentUserId = result.user_id;
                    // 启动心跳
                    heartbeatTimer = setInterval(async function() {
                        try {
                            await fetch("/api/qrcode/heartbeat", {
                                method: "POST",
                                headers: { "Content-Type": "application/json" },
                                body: JSON.stringify({ user_id: currentUserId })
                            });
                        } catch (e) {}
                    }, 2000);
                    // 启动消息轮询
                    messageTimer = setInterval(fetchMessages, 3000);
                    fetchMessages();
                } else {
                    resultDiv.className = "error";
                    resultDiv.textContent = (result.message || "连接失败");
                    btn.textContent = "确认连接";
                    btn.disabled = false;
                }
            } catch (e) {
                console.error('连接失败:', e);
                const resultDiv = document.getElementById("result");
                resultDiv.style.display = "block";
                resultDiv.className = "error";
                resultDiv.textContent = "网络错误，请重试: " + (e.message || e);
                btn.textContent = "确认连接";
                btn.disabled = false;
            }
        }
        
        async function fetchMessages() {
            if (!currentUserId) return;
            try {
                const response = await fetch("/api/phone/messages?user_id=" + encodeURIComponent(currentUserId));
                const result = await response.json();
                console.log('轮询消息结果:', result);
                if (result.success && result.has_new && result.messages.length > 0) {
                    hasNewMessage = true;
                    document.getElementById("newBadge").style.display = "inline-block";
                    result.messages.forEach(msg => {
                        renderMessage(msg);
                    });
                }
            } catch (e) {
                console.error('获取消息失败:', e);
            }
        }
        
        function isImageContent(text) {
            var trimmed = text.trim();
            return /^https?:\\/\\/[^\\s<>]+\\.(?:png|jpg|jpeg|gif|webp)(\\?.*)?$/i.test(trimmed) ||
                   /^data:image\\/(?:png|jpg|jpeg|gif|webp);base64,/i.test(trimmed);
        }
        
        function renderMessage(msg) {
            console.log('渲染消息:', msg);
            const list = document.getElementById("messageList");
            const card = document.createElement("div");
            card.className = "message-card";
            
            // 兼容两种消息结构：formatted (msg.content.content) 和 document (msg.texts/msg.images)
            let texts = [];
            let images = [];
            let isDocument = false;
            let docFilename = '';
            
            if (msg.type === 'document_with_images') {
                isDocument = true;
                docFilename = msg.filename || '';
                texts = msg.texts || [];
                images = msg.images || [];
            } else {
                const content = msg.content || {};
                const innerContent = content.content || content;
                texts = innerContent.texts || [];
                images = innerContent.images || [];
            }
            
            let html = '<div class="message-title">电脑发送的内容</div>';
            
            if (isDocument && docFilename) {
                html += '<div class="message-text" style="background:#EFF6FF;"><a href="/mobile/download/' + encodeURIComponent(docFilename) + '" download style="color:#1E40AF; text-decoration:underline; font-weight:bold;">📄 PDF文档: ' + escapeHtml(docFilename) + '（点击下载）</a></div>';
            }
            
            if (texts && texts.length > 0) {
                texts.forEach(function(text) {
                    if (isImageContent(text)) {
                        html += '<img src="' + escapeHtml(text) + '" class="message-image" alt="图片">';
                    } else {
                        html += '<div class="message-text">' + escapeHtml(text) + '</div>';
                    }
                });
            } else if (!isDocument) {
                html += '<div class="message-text" style="color:#999;">（无文字内容）</div>';
            }
            
            if (images && images.length > 0) {
                images.forEach(function(img, idx) {
                    if (img.base64) {
                        html += '<img src="' + img.base64 + '" class="message-image" alt="' + escapeHtml(img.filename || '图片') + '" onclick="window.open(this.src, \\'_blank\\')" style="cursor:pointer;">';
                        html += '<button class="btn-download" onclick="downloadImage(\\'' + img.base64 + '\\', \\'' + (img.filename || 'image_' + idx + '.png') + '\\')">下载图片</button>';
                    } else if (img.url) {
                        html += '<a href="' + escapeHtml(img.url) + '" target="_blank" style="display:block; margin-bottom:8px;"><img src="' + escapeHtml(img.url) + '" class="message-image" alt="' + escapeHtml(img.filename || '图片') + '"></a>';
                        html += '<a href="' + escapeHtml(img.url) + '" target="_blank" class="btn-download" style="display:inline-block; text-align:center; text-decoration:none;">查看原图</a>';
                    } else if (img.path) {
                        var dlUrl = '/mobile/download/' + encodeURIComponent(img.path);
                        html += '<a href="' + dlUrl + '" target="_blank" style="display:block; margin-bottom:8px;"><img src="' + dlUrl + '" class="message-image" alt="' + escapeHtml(img.filename || '图片') + '"></a>';
                        html += '<a href="' + dlUrl + '" download class="btn-download" style="display:inline-block; text-align:center; text-decoration:none;">下载图片</a>';
                    } else {
                        html += '<div class="message-text">[图片] ' + escapeHtml(img.filename || '未知图片') + ' (' + escapeHtml(img.format || '未知格式') + ')</div>';
                    }
                });
            }
            
            const timeStr = msg.timestamp ? new Date(msg.timestamp).toLocaleString('zh-CN') : new Date().toLocaleString('zh-CN');
            html += '<div class="message-time">' + timeStr + '</div>';
            
            card.innerHTML = html;
            list.insertBefore(card, list.firstChild);
        }
        
        function escapeHtml(text) {
            if (!text) return '';
            const div = document.createElement('div');
            div.textContent = text;
            return div.innerHTML;
        }
        
        function downloadImage(base64Data, filename) {
            const link = document.createElement('a');
            link.href = base64Data;
            link.download = filename;
            document.body.appendChild(link);
            link.click();
            document.body.removeChild(link);
        }
        
        async function disconnect() {
            if (!currentUserId) return;
            try {
                const response = await fetch("/api/qrcode/disconnect", {
                    method: "POST",
                    headers: { "Content-Type": "application/json" },
                    body: JSON.stringify({ user_id: currentUserId })
                });
                const result = await response.json();
                if (result.success) {
                    clearInterval(heartbeatTimer);
                    clearInterval(messageTimer);
                    currentUserId = null;
                    document.getElementById("connectedView").style.display = "none";
                    document.getElementById("messageArea").style.display = "none";
                    document.getElementById("inputForm").style.display = "block";
                    document.getElementById("status").textContent = "已断开连接";
                    document.getElementById("status").className = "status disconnected";
                    const resultDiv = document.getElementById("result");
                    resultDiv.style.display = "block";
                    resultDiv.className = "success";
                    resultDiv.textContent = "已断开连接";
                    document.getElementById("connectBtn").textContent = "重新连接";
                    document.getElementById("connectBtn").disabled = false;
                }
            } catch (e) {
                alert("断开连接失败，请刷新页面重试");
            }
        }
    </script>
</body>
</html>
    '''

# ==================== 手机扫码上传文件到电脑 ====================

UPLOAD_TEMP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads', 'qr_upload')
os.makedirs(UPLOAD_TEMP_DIR, exist_ok=True)
latest_upload = {'filename': None, 'timestamp': 0}

@app.route('/mobile/upload', methods=['GET', 'POST'])
def mobile_upload():
    """手机扫码上传文件页面"""
    if request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有文件'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '文件名为空'}), 400
        try:
            filename = f"{int(__import__('time').time())}_{file.filename}"
            filepath = os.path.join(UPLOAD_TEMP_DIR, filename)
            file.save(filepath)
            global latest_upload
            latest_upload = {'filename': file.filename, 'timestamp': int(__import__('time').time())}
            return jsonify({'success': True, 'message': '上传成功', 'filename': file.filename})
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500
    
    # GET 返回上传页面
    return '''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>上传文件</title>
    <style>
        * { box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; width: 100%; min-height: 100vh; margin: 0; padding: 20px; background: #fff; color: #000; }
        .card { background: #f9f9f9; border: 1px solid #e5e5e5; border-radius: 12px; padding: 24px; }
        h2 { text-align: center; color: #000; margin-bottom: 24px; }
        input[type="file"] { display: none; }
        .file-btn { width: 100%; padding: 14px; background: #fff; color: #000; border: 1px solid #000; border-radius: 8px; font-size: 16px; cursor: pointer; text-align: center; display: block; margin-bottom: 12px; word-break: break-all; overflow-wrap: break-word; }
        .file-btn:active { background: #eee; }
        #result { margin-top: 16px; padding: 12px; border-radius: 8px; text-align: center; display: none; word-break: break-all; overflow-wrap: break-word; }
        .success { background: #f0f0f0; color: #000; border: 1px solid #ccc; }
        .error { background: #f0f0f0; color: #cc0000; border: 1px solid #ccc; }
    </style>
</head>
<body>
    <div class="card">
        <h2>上传文件到电脑</h2>
        <label class="file-btn">
            <input type="file" id="fileInput" onchange="uploadFile()">
            选择文件
        </label>
        <div id="result"></div>
    </div>
    <script>
        async function uploadFile() {
            const input = document.getElementById("fileInput");
            const file = input.files[0];
            if (!file) return;
            const formData = new FormData();
            formData.append("file", file);
            const resultDiv = document.getElementById("result");
            resultDiv.style.display = "block";
            resultDiv.className = "";
            resultDiv.textContent = "上传中...";
            try {
                const response = await fetch("/mobile/upload", {
                    method: "POST",
                    body: formData
                });
                const result = await response.json();
                if (result.success) {
                    resultDiv.className = "success";
                    resultDiv.textContent = "✓ 上传成功: " + result.filename;
                } else {
                    resultDiv.className = "error";
                    resultDiv.textContent = "✗ " + (result.message || "上传失败");
                }
            } catch (e) {
                resultDiv.className = "error";
                resultDiv.textContent = "✗ 网络错误";
            }
        }
    </script>
</body>
</html>
    '''

@app.route('/api/upload/status', methods=['GET'])
def api_upload_status():
    """查询是否有新上传的文件"""
    since = request.args.get('since', 0, type=int)
    if latest_upload['filename'] and latest_upload['timestamp'] > since:
        return jsonify({
            'success': True,
            'has_new': True,
            'filename': latest_upload['filename'],
            'timestamp': latest_upload['timestamp']
        })
    return jsonify({
        'success': True,
        'has_new': False
    })

@app.route('/mobile/download/<path:filename>')
def mobile_download(filename):
    """提供PDF文档和图片文件的下载"""
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'old', 'phone', 'output')
    file_path = os.path.join(output_dir, filename)
    if not os.path.exists(file_path) or not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'}), 404
    return send_from_directory(output_dir, filename, as_attachment=False)

@app.route('/api/chat/langchain', methods=['POST'])
def api_chat_langchain():
    try:
        from ruoyi_langchain.chat_service import ChatService
        
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': '请求体不能为空'
            }), 400
        
        message = data.get('message', '')
        if not message or not isinstance(message, str):
            return jsonify({
                'success': False,
                'message': '消息内容不能为空'
            }), 400
        
        model = data.get('model', 'qwen-turbo')
        if not isinstance(model, str):
            return jsonify({
                'success': False,
                'message': '模型参数必须是字符串'
            }), 400
        
        conversation_id = data.get('conversation_id')
        if conversation_id and not isinstance(conversation_id, str):
            return jsonify({
                'success': False,
                'message': '会话ID必须是字符串'
            }), 400
        
        agent_type = data.get('agent_type', 'personal')
        if not isinstance(agent_type, str):
            return jsonify({
                'success': False,
                'message': 'Agent类型必须是字符串'
            }), 400
        
        chat_service = ChatService()
        
        if '@知识库' in message:
            result = chat_service.chat(
                message=message,
                agent_type='knowledge',
                model=model,
                conversation_id=conversation_id
            )
        elif '@组织管理' in message:
            result = chat_service.chat(
                message=message,
                agent_type='knowledge',
                model=model,
                conversation_id=conversation_id
            )
        elif '@联网搜索' in message:
            result = chat_service.chat(
                message=message,
                agent_type='search',
                model=model,
                conversation_id=conversation_id
            )
        else:
            result = chat_service.chat(
                message=message,
                agent_type=agent_type,
                model=model,
                conversation_id=conversation_id
            )
        
        return jsonify(result)
    except Exception as e:
        print(f'[api_chat_langchain] 错误: {e}', flush=True)
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/vector/search', methods=['POST'])
def api_vector_search():
    try:
        from vector_store.chroma_store import CharmDBStore
        
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': '请求体不能为空'
            }), 400
        
        query = data.get('query', '')
        if not query or not isinstance(query, str):
            return jsonify({
                'success': False,
                'message': '查询关键词不能为空'
            }), 400
        
        k = data.get('k', 5)
        if not isinstance(k, int) or k <= 0 or k > 100:
            return jsonify({
                'success': False,
                'message': 'k值必须是1-100之间的整数'
            }), 400
        
        collection_name = data.get('collection_name', 'knowledge_base')
        if not isinstance(collection_name, str):
            return jsonify({
                'success': False,
                'message': '集合名称必须是字符串'
            }), 400
        
        chroma_store = CharmDBStore(persist_directory=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chroma_db'))
        results = chroma_store.similarity_search(collection_name, query, k=k)
        
        response_data = []
        for doc in results:
            response_data.append({
                'content': doc.page_content,
                'metadata': doc.metadata
            })
        
        return jsonify({
            'success': True,
            'data': response_data
        })
    except Exception as e:
        print(f'[api_vector_search] 错误: {e}', flush=True)
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/vector/collections', methods=['GET'])
def api_vector_collections():
    try:
        from vector_store.chroma_store import CharmDBStore
        
        chroma_store = CharmDBStore(persist_directory=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chroma_db'))
        collections = chroma_store.get_all_collections()
        
        stats = []
        for col_name in collections:
            try:
                stats.append(chroma_store.get_collection_stats(col_name))
            except:
                pass
        
        return jsonify({
            'success': True,
            'data': stats
        })
    except Exception as e:
        print(f'[api_vector_collections] 错误: {e}', flush=True)
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

if __name__ == '__main__':
    print("Starting Flask development server on port 5000...")
    print("Image upload directory:", IMAGE_DIR)
    app.run(host='0.0.0.0', port=5000, debug=False)
